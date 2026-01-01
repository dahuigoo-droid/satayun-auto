# -*- coding: utf-8 -*-
"""
🤖 Claude API 연동 모듈
- 장별 해석 자동 생성
- 프롬프트 관리
- Docx 출력
"""

import os
from typing import Optional, List, Dict, Callable
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


# ============================================
# Claude API 클라이언트
# ============================================
class SajuInterpreter:
    """사주 해석 생성기"""
    
    def __init__(self, api_key: str, model: str = "claude-sonnet-4-20250514"):
        """
        Args:
            api_key: Anthropic API 키
            model: 사용할 모델
                - claude-sonnet-4-20250514: 빠르고 경제적 (추천)
                - claude-haiku-3-5-20241022: 가장 저렴
        """
        import anthropic
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = model
    
    def generate_chapter(
        self,
        chapter_name: str,
        gpt_text: str,
        prompt_template: str,
        max_tokens: int = 4000
    ) -> str:
        """
        단일 장 해석 생성
        
        Args:
            chapter_name: 장 이름 (예: "원국분석")
            gpt_text: 사주 데이터 텍스트
            prompt_template: 프롬프트 템플릿
            max_tokens: 최대 토큰 수
            
        Returns:
            생성된 해석 텍스트
        """
        full_prompt = f"""{prompt_template}

[사주 데이터]
{gpt_text}

위 데이터를 바탕으로 "{chapter_name}" 장을 작성해주세요."""
        
        try:
            response = self.client.messages.create(
                model=self.model,
                max_tokens=max_tokens,
                messages=[
                    {"role": "user", "content": full_prompt}
                ]
            )
            return response.content[0].text
        except Exception as e:
            return f"[오류] {chapter_name} 생성 실패: {str(e)}"
    
    def generate_all_chapters(
        self,
        gpt_text: str,
        prompts: List[Dict],  # [{"name": str, "template": str}, ...]
        progress_callback: Optional[Callable[[str, int, int], None]] = None
    ) -> Dict[str, str]:
        """
        모든 장 해석 일괄 생성
        
        Args:
            gpt_text: 사주 데이터 텍스트
            prompts: [{"name": 장이름, "template": 프롬프트}, ...]
            progress_callback: (장이름, 현재, 전체) 콜백
            
        Returns:
            {장이름: 해석텍스트} 딕셔너리
        """
        results = {}
        total = len(prompts)
        
        for idx, prompt_info in enumerate(prompts):
            chapter_name = prompt_info["name"]
            template = prompt_info["template"]
            
            if progress_callback:
                progress_callback(chapter_name, idx + 1, total)
            
            results[chapter_name] = self.generate_chapter(
                chapter_name=chapter_name,
                gpt_text=gpt_text,
                prompt_template=template
            )
        
        return results


# ============================================
# 프롬프트 관리
# ============================================
def load_prompts_from_dir(prompts_dir: str = "prompts") -> List[Dict]:
    """
    프롬프트 폴더에서 모든 프롬프트 로드
    
    Returns:
        [{"num": str, "name": str, "template": str, "path": str}, ...]
    """
    prompts = []
    
    if not os.path.exists(prompts_dir):
        return prompts
    
    for filename in sorted(os.listdir(prompts_dir)):
        if filename.endswith('.txt'):
            parts = filename.replace('.txt', '').split('_', 1)
            if len(parts) == 2:
                num, name = parts
            else:
                num, name = "00", parts[0]
            
            path = os.path.join(prompts_dir, filename)
            with open(path, 'r', encoding='utf-8') as f:
                template = f.read()
            
            prompts.append({
                "num": num,
                "name": name,
                "template": template,
                "path": path
            })
    
    return prompts


def list_prompts(prompts_dir: str = "prompts") -> List[tuple]:
    """
    사용 가능한 프롬프트 목록
    
    Returns:
        [(번호, 장이름, 파일경로), ...]
    """
    prompts = load_prompts_from_dir(prompts_dir)
    return [(p["num"], p["name"], p["path"]) for p in prompts]


# ============================================
# Docx 생성
# ============================================
def create_chapter_docx(
    chapter_num: int,
    chapter_name: str,
    content: str,
    output_path: str,
    image_tag: str = None,
    font_name: str = "맑은 고딕"
) -> str:
    """
    단일 장 Docx 파일 생성
    
    Args:
        chapter_num: 장 번호
        chapter_name: 장 이름
        content: 해석 텍스트
        output_path: 저장 경로
        image_tag: 이미지 태그 (예: "01_원국표")
        font_name: 폰트 이름
        
    Returns:
        저장된 파일 경로
    """
    doc = Document()
    
    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = font_name
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    # 장 제목
    title = f"제{chapter_num}장 {chapter_name}"
    doc.add_paragraph(title)
    
    # 이미지 태그 (있으면)
    if image_tag:
        doc.add_paragraph(f"{{{{IMG:{image_tag}}}}}")
    
    # 본문
    paragraphs = content.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            doc.add_paragraph(para_text.strip())
    
    doc.save(output_path)
    return output_path


def create_all_chapter_docx(
    chapters: Dict[str, str],
    output_dir: str,
    customer_name: str,
    image_mapping: Dict[str, str] = None
) -> List[str]:
    """
    모든 장 Docx 파일 일괄 생성
    
    Args:
        chapters: {장이름: 해석텍스트}
        output_dir: 출력 폴더
        customer_name: 고객명
        image_mapping: {장이름: 이미지태그}
        
    Returns:
        생성된 파일 경로 리스트
    """
    os.makedirs(output_dir, exist_ok=True)
    
    if image_mapping is None:
        image_mapping = {
            '원국분석': '01_원국표',
            '대운분석': '02_대운표',
            '세운분석': '03_세운표',
            '월운분석': '04_월운표',
            '용신분석': '16_용신표',
            '오행분석': '05_오행분석',
        }
    
    paths = []
    for idx, (name, content) in enumerate(chapters.items(), 1):
        filename = f"제{idx}장_{name}.docx"
        output_path = os.path.join(output_dir, filename)
        
        image_tag = image_mapping.get(name)
        
        create_chapter_docx(
            chapter_num=idx,
            chapter_name=name,
            content=content,
            output_path=output_path,
            image_tag=image_tag
        )
        
        paths.append(output_path)
    
    return paths


# ============================================
# 비용 계산
# ============================================
def estimate_cost(
    num_chapters: int,
    avg_input_tokens: int = 3000,
    avg_output_tokens: int = 2000,
    model: str = "claude-sonnet-4-20250514"
) -> Dict:
    """
    예상 비용 계산
    
    Returns:
        {'input_tokens', 'output_tokens', 'cost_usd', 'cost_krw'}
    """
    pricing = {
        "claude-sonnet-4-20250514": {"input": 3.0, "output": 15.0},
        "claude-haiku-3-5-20241022": {"input": 0.25, "output": 1.25},
    }
    
    if model not in pricing:
        model = "claude-sonnet-4-20250514"
    
    total_input = num_chapters * avg_input_tokens
    total_output = num_chapters * avg_output_tokens
    
    cost_usd = (
        (total_input / 1_000_000) * pricing[model]["input"] +
        (total_output / 1_000_000) * pricing[model]["output"]
    )
    
    return {
        "input_tokens": total_input,
        "output_tokens": total_output,
        "cost_usd": round(cost_usd, 4),
        "cost_krw": int(cost_usd * 1400)
    }


# ============================================
# 테스트
# ============================================
if __name__ == "__main__":
    print("Claude API 모듈 로드 완료")
    
    # 비용 예측 테스트
    cost = estimate_cost(num_chapters=10)
    print(f"10장 생성 예상 비용: ${cost['cost_usd']} (약 {cost['cost_krw']}원)")
